"""
Document processing service for text extraction and thumbnail generation.

This module provides utilities for processing document files including:
- Text extraction from PDF, DOCX, and Markdown files
- Thumbnail generation from PDF first pages
"""

import logging
import os
import tempfile
from typing import Optional, Tuple
from io import BytesIO

from PIL import Image
from PyPDF2 import PdfReader

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing document files."""

    def __init__(self):
        """Initialize document processor."""
        self.thumbnail_size = 300
        self.thumbnail_quality = 85

    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text as string
        """
        try:
            reader = PdfReader(file_path)
            text_parts = []

            for page in reader.pages:
                try:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from PDF page: {e}")
                    continue

            extracted_text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(extracted_text)} characters from PDF")
            return extracted_text

        except Exception as e:
            logger.error(f"Failed to extract text from PDF {file_path}: {e}")
            return ""

    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text as string
        """
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            extracted_text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(extracted_text)} characters from DOCX")
            return extracted_text

        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            return ""

    def extract_text_from_markdown(self, file_path: str) -> str:
        """
        Extract text from Markdown file.

        Args:
            file_path: Path to Markdown file

        Returns:
            Extracted text as string (raw markdown content)
        """
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

            logger.info(f"Extracted {len(text)} characters from Markdown")
            return text

        except Exception as e:
            logger.error(f"Failed to extract text from Markdown {file_path}: {e}")
            return ""

    def generate_pdf_thumbnail(
        self, file_path: str, size: Optional[int] = None
    ) -> Optional[bytes]:
        """
        Generate thumbnail from first page of PDF.

        Args:
            file_path: Path to PDF file
            size: Thumbnail size (square), defaults to self.thumbnail_size

        Returns:
            Thumbnail image as bytes (JPEG format), or None if failed
        """
        if size is None:
            size = self.thumbnail_size

        try:
            from pdf2image import convert_from_path

            # Convert only the first page
            images = convert_from_path(
                file_path, first_page=1, last_page=1, size=(size, None)
            )

            if not images:
                logger.warning(f"No images generated from PDF {file_path}")
                return None

            # Get the first page image
            first_page = images[0]

            # Convert to RGB if needed (for JPEG)
            if first_page.mode != "RGB":
                first_page = first_page.convert("RGB")

            # Resize to fit within square thumbnail while maintaining aspect ratio
            first_page.thumbnail((size, size), Image.Resampling.LANCZOS)

            # Save to bytes
            thumbnail_io = BytesIO()
            first_page.save(
                thumbnail_io, format="JPEG", quality=self.thumbnail_quality, optimize=True
            )
            thumbnail_bytes = thumbnail_io.getvalue()

            logger.info(f"Generated PDF thumbnail: {len(thumbnail_bytes)} bytes")
            return thumbnail_bytes

        except Exception as e:
            logger.error(f"Failed to generate PDF thumbnail for {file_path}: {e}")
            return None

    def generate_document_thumbnail(
        self, file_path: str, content_type: str, size: Optional[int] = None
    ) -> Optional[bytes]:
        """
        Generate thumbnail for document based on file type.

        Currently only supports PDF thumbnails. Other document types
        will return None (can use generic icons on frontend).

        Args:
            file_path: Path to document file
            content_type: MIME type of document
            size: Thumbnail size (square)

        Returns:
            Thumbnail bytes or None
        """
        if content_type == "application/pdf":
            return self.generate_pdf_thumbnail(file_path, size)

        # For other document types (DOCX, MD), no thumbnail generation yet
        # Frontend can display generic document icons
        logger.info(f"No thumbnail generation for content type: {content_type}")
        return None

    def extract_text(self, file_path: str, content_type: str) -> str:
        """
        Extract text from document based on file type.

        Args:
            file_path: Path to document file
            content_type: MIME type of document

        Returns:
            Extracted text string
        """
        if content_type == "application/pdf":
            return self.extract_text_from_pdf(file_path)

        elif content_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ]:
            return self.extract_text_from_docx(file_path)

        elif content_type in ["text/markdown", "text/x-markdown"]:
            return self.extract_text_from_markdown(file_path)

        else:
            logger.warning(f"No text extraction support for content type: {content_type}")
            return ""


# Singleton instance
document_processor = DocumentProcessor()





